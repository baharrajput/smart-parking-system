from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(3.17)
section.right_margin  = Cm(2.54)

def set_run(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run(run, size=14, bold=True)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run(run, size=12, bold=True, italic=True)
    return p

def para(text, first_line=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run(run, size=12)
    return p

def para_no_indent(text):
    return para(text, first_line=False)

def bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(u"•  " + text)
    set_run(run, size=12)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EFEFEF')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(text)
    set_run(run, name="Courier New", size=10)
    return p

def simple_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        r = c.paragraphs[0].runs[0]
        set_run(r, size=11, bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        c._tc.get_or_add_tcPr().append(shd)
    for ri, rd in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(rd):
            c = row.cells[ci]
            c.text = str(val)
            set_run(c.paragraphs[0].runs[0], size=11)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Smart Parking Detection System"), size=22, bold=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("A Final Project Report"), size=14, italic=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Submitted in Partial Fulfillment of the Requirements\nfor the Course of Computer Vision"), size=12)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Submitted by\n"), size=12)
set_run(p.add_run("Hassan Bahar"), size=13, bold=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run(datetime.date.today().strftime("%B %Y")), size=12)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────
h1("Abstract")
para(
    "This report describes a parking slot occupancy detection system built using Python and "
    "OpenCV. The goal was to create a working prototype that can monitor a physical parking "
    "layout in real time through a mobile phone camera and indicate which slots are free and "
    "which are taken. The system uses a combination of image differencing, edge analysis, and "
    "background subtraction to make this determination. A one-time calibration step captures "
    "the appearance of each empty slot, and from that point on the system runs without any "
    "manual input."
)

# ─────────────────────────────────────────────────────────────────────────────
# 1. INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
h1("1.  Introduction")
para(
    "Finding an available parking spot in a busy area is something most people deal with on a "
    "daily basis. In larger parking facilities this often means driving through each row and "
    "checking manually, which wastes both time and fuel. An automated system that can tell "
    "drivers which spots are free before they enter would be genuinely useful, and this project "
    "is an attempt to build a small-scale version of that idea."
)
para(
    "The setup for this project is intentionally low-cost. Instead of mounting dedicated "
    "sensors on each parking slot or installing an expensive IP camera system, the entire "
    "input comes from an ordinary Android phone running the IP Webcam application. The phone "
    "is placed overhead and streams live video over the local WiFi network to a laptop where "
    "the detection runs. The parking layout itself is a hand-drawn grid on paper with 16 "
    "labeled spots arranged in four columns of four."
)
para(
    "The technical challenge is to reliably decide, for each of these 16 rectangles in the "
    "video frame, whether something is sitting on it or not. This turns out to be harder than "
    "it sounds because the paper itself has drawn lines and labels that create visual noise, "
    "and the camera image quality and lighting can change between sessions. The approach taken "
    "here combines several classical computer vision techniques so that no single method needs "
    "to be perfect on its own."
)

# ─────────────────────────────────────────────────────────────────────────────
# 2. SYSTEM OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────
h1("2.  System Overview")
para(
    "At the highest level the system has three phases: a one-time slot definition step, "
    "an automatic startup calibration, and a continuous detection loop."
)
para(
    "The slot definition is done once by running a setup tool that opens the camera feed "
    "on screen. The user clicks the top-left corner of the paper layout and then the "
    "bottom-right corner. The software divides that rectangle into a 4x4 grid and saves "
    "the pixel coordinates of all 16 slots to a JSON file on disk. This file persists "
    "between sessions so the step never needs to be repeated as long as the camera stays "
    "in the same position."
)
para(
    "When the main program starts, it reads those saved coordinates and connects to the "
    "phone camera. It then waits five seconds while the camera exposure stabilizes and "
    "captures a reference image of each empty slot. These reference images are stored in "
    "memory as the baseline for what each slot looks like when nothing is parked there."
)
para(
    "After that the system enters the detection loop. On every frame it crops each slot "
    "out of the image, runs a set of image analysis operations, and produces a score "
    "between 0 and 1 indicating how different that slot looks from its empty reference. "
    "If the score is above a threshold the slot is marked occupied and highlighted in red; "
    "otherwise it is shown in green. The results are drawn onto the screen in real time "
    "alongside a panel showing counts and row-level breakdowns."
)

# ─────────────────────────────────────────────────────────────────────────────
# 3. HARDWARE AND SOFTWARE USED
# ─────────────────────────────────────────────────────────────────────────────
h1("3.  Hardware and Software Used")

h2("Hardware")
para_no_indent(
    "The only dedicated hardware beyond a laptop is an Android mobile phone. The IP Webcam "
    "application (free, available on Google Play) turns the phone into a network camera that "
    "streams MJPEG video over HTTP. Both the phone and laptop need to be on the same WiFi "
    "network. The phone is positioned overhead pointing down at the paper layout."
)

h2("Software and Libraries")
para_no_indent(
    "The entire implementation is in Python 3. The main dependency is OpenCV (cv2 version 4), "
    "which handles camera capture, all image processing operations, and drawing the interface. "
    "NumPy is used for array math. No deep learning framework or pre-trained model is involved; "
    "everything is classical computer vision."
)
simple_table(
    ["Library / Tool", "What it is used for"],
    [
        ["OpenCV (cv2)", "Camera capture, image processing, drawing the UI"],
        ["NumPy", "Array operations and pixel-level math"],
        ["IP Webcam (Android app)", "Streaming live video from the phone over WiFi"],
        ["JSON (Python stdlib)", "Saving and loading slot coordinates between sessions"],
        ["Python 3 / Windows", "Runtime environment"],
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 4. COMPUTER VISION TECHNIQUES
# ─────────────────────────────────────────────────────────────────────────────
h1("4.  Computer Vision Techniques")
para(
    "Each parking slot is analysed independently. The full frame is captured from the camera, "
    "each slot's region of interest (ROI) is cropped from it, and a pipeline of operations is "
    "applied to that small patch. The following subsections describe each technique in the order "
    "it appears in the pipeline."
)

h2("4.1  Grayscale Conversion and Noise Reduction")
para_no_indent(
    "The cropped ROI is first converted from BGR colour to grayscale using "
    "cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY). Colour information is not needed for occupancy "
    "detection and removing it reduces computation. The grayscale image then goes through "
    "two smoothing steps. A median filter (cv2.medianBlur with kernel size 5) removes "
    "salt-and-pepper noise that appears in JPEG-compressed camera streams. This is followed "
    "by a Gaussian blur (cv2.GaussianBlur with a 5x5 kernel) which smooths out remaining "
    "high-frequency variation before thresholding."
)

h2("4.2  Adaptive Thresholding")
para_no_indent(
    "Converting the blurred grayscale image to a binary mask is done with adaptive thresholding "
    "rather than a fixed global threshold. With a global threshold, areas that are in shadow "
    "would be incorrectly classified because the same absolute value does not mean the same "
    "thing across different parts of the image. Adaptive thresholding computes a local "
    "threshold for each pixel based on the weighted mean intensity of its surrounding "
    "neighbourhood and subtracts a small constant from it. This makes the method self-adjusting "
    "across different lighting conditions. The result is inverted so that dark regions, "
    "which correspond to objects, appear as white in the binary mask. A neighbourhood block "
    "size of 21 pixels and a subtraction constant of 12 were found to give clean results "
    "on this particular setup through experimentation."
)

h2("4.3  Morphological Operations")
para_no_indent(
    "The binary mask produced by thresholding still contains small scattered noise pixels "
    "and may have gaps inside solid objects. Two morphological operations clean this up. "
    "An opening operation (erosion followed by dilation) removes isolated noise pixels "
    "that are smaller than the structuring element. A closing operation (dilation followed "
    "by erosion) fills small holes inside detected regions. Both operations use an elliptical "
    "structuring element of size 5x5 pixels. The opening is applied once and the closing "
    "twice, which was found to give the cleanest binary masks without eroding genuine "
    "object boundaries."
)

h2("4.4  Reference Frame Comparison (Primary Method)")
para_no_indent(
    "The most important technique in the pipeline is comparing the current slot ROI against "
    "the reference image taken when that slot was known to be empty. This is done with "
    "cv2.absdiff, which computes the absolute per-pixel difference between two images. "
    "The mean of this difference, normalised to the range 0 to 1 by dividing by 255, gives "
    "a score that directly measures how much the slot has changed from its empty state. "
    "A score near zero means the slot looks the same as when it was empty; a higher score "
    "indicates that something has been added."
)
para_no_indent(
    "This approach sidesteps most of the problems with detecting on a visually complex "
    "background like a hand-drawn paper layout. Since the reference was captured with all "
    "those lines and labels already present, they contribute equally to both images and "
    "cancel out in the subtraction. Only things that have genuinely changed — an object "
    "placed on the slot — produce a non-zero difference."
)

h2("4.5  Edge Density via Canny Edge Detection")
para_no_indent(
    "As a supporting signal, the edge density of each slot ROI is computed using the Canny "
    "edge detector. An occupied slot with an object in it tends to have more internal edges "
    "than an empty slot. Canny works by computing image gradients with Sobel operators, "
    "suppressing non-maximal gradient responses to produce thin edges, and then applying "
    "hysteresis thresholding with a lower bound of 40 and upper bound of 120 to decide "
    "which edges are real. The edge score is the fraction of edge pixels in the slot "
    "region, giving a value between 0 and 1."
)

h2("4.6  Background Subtraction with MOG2")
para_no_indent(
    "OpenCV's MOG2 background subtractor models the background of each pixel as a mixture "
    "of Gaussians and classifies any pixel that does not fit the model as foreground. It "
    "has built-in shadow detection, which is enabled here so that shadows cast by objects "
    "are not mistaken for the objects themselves. The foreground mask is thresholded at 200 "
    "to remove shadow pixels and keep only definite foreground. The subtractor is configured "
    "with a history of 300 frames and a variance threshold of 40, with shadow detection enabled. "
    "Shadow pixels, which MOG2 marks with a value of 127, are removed by applying a threshold "
    "of 200 to the foreground mask so only fully detected foreground remains."
)
para_no_indent(
    "MOG2 works best after seeing several hundred frames of background, so it is used as "
    "a minor supporting signal rather than the primary detector. Its weight in the final "
    "formula is kept low at 5%."
)

h2("4.7  Combined Score and Decision")
para_no_indent(
    "The individual signals are combined into a single weighted score. The weights reflect "
    "how reliable each signal is for this specific application. Reference frame comparison "
    "gets the largest weight because it directly measures change from the known empty state. "
    "When no reference is available, which only happens during the very first frames before "
    "calibration completes, a fallback formula is used that distributes the weight across "
    "the remaining signals: 40% pixel density, 30% edge density, 20% contour area, and "
    "10% background subtraction. In normal operation with a reference frame, the weights "
    "are 50% for frame difference, 20% for pixel density, 15% for edge density, 10% for "
    "contour analysis, and 5% for background subtraction. A slot is classified as occupied "
    "when the final score exceeds 0.15, a threshold that can be adjusted at runtime using "
    "the plus and minus keys."
)

# ─────────────────────────────────────────────────────────────────────────────
# 5. DATA STORAGE
# ─────────────────────────────────────────────────────────────────────────────
h1("5.  Data Storage")
para(
    "Two files on disk carry state between sessions. The first is data/slots.json, which "
    "stores the pixel bounding box of each of the 16 parking slots as a JSON array. This "
    "is written once during the setup step and read back at the start of every subsequent "
    "run. Each entry in the array records the slot number, the column group it belongs to "
    "(labelled as row 1 through 4), and four pixel coordinates giving the top-left and "
    "bottom-right corners of the slot rectangle within the 1280 by 720 frame."
)
para(
    "The second file is data/reference_frame.jpg, saved at the same time as slots.json. "
    "This is a snapshot of the camera view taken immediately after the user finishes clicking "
    "the corners during setup. It serves as a visual record but is not loaded at runtime — "
    "the in-memory reference grays computed during startup calibration are used instead, "
    "which are always freshly captured from the live camera on each run."
)
para(
    "All configurable parameters live in config.py — the camera URL, frame resolution, "
    "detection threshold, filter kernel sizes, and file paths. Keeping these in one place "
    "makes it straightforward to adjust the system without editing the processing code."
)

# ─────────────────────────────────────────────────────────────────────────────
# 6. CODE STRUCTURE
# ─────────────────────────────────────────────────────────────────────────────
h1("6.  Code Structure")
para(
    "The project is split into a small set of focused files rather than putting everything "
    "in one large script. This made development and debugging easier because each part could "
    "be tested independently."
)
doc.add_picture(
    r"d:\Spring 2026\Computer Vision\Final_Project\file_tree.png",
    width=Inches(5.8)
)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
para(
    "The SlotDetector class in core/detector.py contains all the image processing described "
    "in section 4. Its public interface is a calibrate() method that saves empty-state "
    "references for all slots, and an analyze_all() method that processes a frame and "
    "returns a list of results. The UI and the main loop do not need to know anything about "
    "the computer vision internals."
)
para(
    "The SidePanelUI class takes the frame and the list of results and produces the final "
    "1600x720 canvas. It draws the slot overlays on the left side and the statistics panel "
    "on the right. All drawing uses cv2 primitives — rectangles, circles, putText, and "
    "addWeighted for transparent fills."
)

# ─────────────────────────────────────────────────────────────────────────────
# 7. INTERFACE
# ─────────────────────────────────────────────────────────────────────────────
h1("7.  Interface")
para(
    "The display window runs fullscreen. The left portion shows the live camera feed with "
    "each slot outlined. A green box means the slot is free; a red box means it is occupied. "
    "Each box has a label in the corner (S1 through S16), the status word FREE or OCC printed "
    "inside it, and a thin progress bar at the bottom showing the detection confidence score."
)
para(
    "The right side is a dark panel about 320 pixels wide. At the top it shows large digit "
    "counters for the total number of free and occupied slots side by side. Below that is "
    "a colour bar that splits the total capacity into green (free) and red (occupied) "
    "proportions. The next section breaks down the status row by row, each row showing "
    "a label, a text summary, and small coloured dots for each individual slot. At the "
    "bottom is a list of all 16 slots with their current state. If every slot is occupied "
    "a banner reading PARKING FULL appears at the very bottom of the panel."
)
para(
    "The keyboard controls are Q to quit, C to recalibrate using the current frame as the "
    "new empty reference, and the plus and minus keys to adjust the detection threshold up "
    "or down by 0.01 per press."
)

# ─────────────────────────────────────────────────────────────────────────────
# 8. RESULTS
# ─────────────────────────────────────────────────────────────────────────────
h1("8.  Results")
para(
    "Testing was done by placing objects (bottles, erasers) on the paper layout while "
    "the system was running and observing the response. After a successful calibration "
    "with the parking empty, placing an object on any slot caused that slot to switch "
    "to red within one or two frames. Removing the object restored it to green within "
    "the same latency. The system ran at roughly 20 frames per second on a standard "
    "laptop connected to the phone over home WiFi."
)
para(
    "The main failure mode encountered during testing was false positives at startup before "
    "calibration completed. This was addressed by extending the warmup period to five seconds "
    "and making sure the reference is captured from a stable frame. Lighting changes between "
    "sessions occasionally caused drift, which the C recalibration key handles cleanly."
)
para(
    "The threshold of 0.15 worked well in practice. Values below about 0.10 produced "
    "occasional false occupancy readings from camera noise; values above 0.20 sometimes "
    "missed small objects. The interactive +/- adjustment makes it easy to tune on the fly "
    "without restarting."
)

# ─────────────────────────────────────────────────────────────────────────────
# 9. DIFFICULTIES ENCOUNTERED
# ─────────────────────────────────────────────────────────────────────────────
h1("9.  Difficulties Encountered")
para(
    "The biggest practical difficulty was that the parking paper itself, with its drawn "
    "boxes and labels, created a visually busy background that confused simple pixel-counting "
    "approaches. The first version of the detector used adaptive thresholding and contour "
    "analysis alone, which reported most slots as occupied even when they were empty because "
    "the drawn lines were being detected as objects. Switching to reference-frame comparison "
    "as the dominant signal fixed this because the lines appear identically in both the "
    "reference and the live frame and cancel out."
)
para(
    "Defining the slot positions was also harder than expected. The initial approach let "
    "the user click and drag each of the 16 slots individually, which was tedious and "
    "prone to misalignment. This was replaced with the two-click corner approach where "
    "the software generates the grid mathematically, which took only a few seconds and "
    "produced perfectly aligned non-overlapping slots."
)
para(
    "The mobile camera IP address changes whenever the phone reconnects to WiFi, requiring "
    "a manual update in config.py. This is a minor but recurring inconvenience. In a "
    "production version this would be handled by either reserving a fixed IP on the router "
    "or using a service discovery mechanism."
)

# ─────────────────────────────────────────────────────────────────────────────
# 10. CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
h1("10.  Conclusion")
para(
    "The system works as intended for the stated goal: given a fixed overhead camera and "
    "a hand-drawn parking layout, it detects which slots are occupied in real time and "
    "displays the result clearly. The combination of reference-frame differencing with "
    "supporting signals from edge density and background subtraction proved more robust "
    "than any single technique would have been on its own."
)
para(
    "The project was useful for understanding how multiple computer vision operations fit "
    "together in a real pipeline. Techniques that seem straightforward in isolation — "
    "thresholding, edge detection, background subtraction — each have failure cases that "
    "need to be compensated for by other signals. The weighted combination approach "
    "provided a practical way to balance them."
)
para(
    "Possible extensions include using a trained object detector like YOLO to replace the "
    "classical pipeline, adding a web interface so the parking status could be viewed "
    "from a phone browser, and logging occupancy over time to generate usage statistics."
)

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
h1("References")
refs = [
    "Bradski, G. (2000). The OpenCV Library. Dr. Dobb's Journal of Software Tools.",
    "OpenCV Documentation. Adaptive Thresholding. https://docs.opencv.org/",
    "OpenCV Documentation. Background Subtraction (MOG2). https://docs.opencv.org/",
    "Canny, J. (1986). A Computational Approach to Edge Detection. "
        "IEEE Transactions on Pattern Analysis and Machine Intelligence, 8(6), 679-698.",
    "Serra, J. (1983). Image Analysis and Mathematical Morphology. Academic Press.",
    "IP Webcam Android Application. Pavel Khlebovich. Google Play Store.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    set_run(p.add_run(f"[{i}]  {r}"), size=11)

out = r"d:\Spring 2026\Computer Vision\Final_Project\Smart_Parking_Report_Final.docx"
doc.save(out)
print("Saved:", out)
